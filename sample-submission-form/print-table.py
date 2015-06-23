from docx import Document

d = Document("/home/paalmbj/sample-submission.docx")

for row in d.tables[0].rows:
    for cell in row.cells:
        print " | ", cell.text
    print " ------- "

for row in d.tables[1].rows:
    for cell in row.cells:
        print "2| ", cell.text
    print "2------- "
